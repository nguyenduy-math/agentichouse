"""DomainAgent — hybrid RAG pipeline scoped to a single domain."""

import asyncio
import io
import unicodedata
import uuid
from pathlib import Path

import pdfplumber
import structlog
from docx import Document

from app.config import settings
from app.services.bm25_service import BM25Service
from app.services.embedding_service import EmbeddingService
from app.services.hybrid_service import rrf_merge
from app.services.rerank_service import RerankService
from app.services.vector_service import VectorService
from app.prompts.chat_prompts import build_system_instruction
from app.utils.text_splitter import split_by_article

log = structlog.get_logger()


def _nfc(text: str) -> str:
    return unicodedata.normalize("NFC", text)


def _parse_file(file_bytes: bytes, filename: str) -> list[dict]:
    ext = Path(filename).suffix.lower()
    pages = []
    if ext == ".pdf":
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = _nfc(page.extract_text() or "")
                if text.strip():
                    pages.append({"text": text, "page_number": i})
    elif ext in (".docx", ".doc"):
        doc = Document(io.BytesIO(file_bytes))
        lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text)
        pages.append({"text": _nfc("\n".join(lines)), "page_number": 1})
    else:
        text = _nfc(file_bytes.decode("utf-8", errors="replace"))
        pages.append({"text": text, "page_number": 1})
    return pages


class DomainAgent:
    """Owns a domain-scoped BM25 index and ChromaDB collection."""

    def __init__(
        self,
        domain_key: str,
        domain_label: str,
        bm25: BM25Service,
        vector: VectorService,
        embedding: EmbeddingService,
        reranker: RerankService,
    ) -> None:
        self.domain_key = domain_key
        self.domain_label = domain_label
        self._bm25 = bm25
        self._vector = vector
        self._embedding = embedding
        self._reranker = reranker

    async def ingest(self, file_bytes: bytes, filename: str) -> list[dict]:
        pages = _parse_file(file_bytes, filename)

        chunks: list[dict] = []
        for page in pages:
            for idx, text in enumerate(split_by_article(page["text"], settings.chunk_size)):
                chunks.append(
                    {
                        "chunk_id": str(uuid.uuid4()),
                        "text": text,
                        "metadata": {
                            "source_file": filename,
                            "page_number": page["page_number"],
                            "chunk_index": idx,
                            "domain": self.domain_key,
                        },
                    }
                )
        if not chunks:
            return []

        texts = [c["text"] for c in chunks]
        embeddings = await self._embedding.embed_documents(texts)
        self._vector.add(
            chunk_ids=[c["chunk_id"] for c in chunks],
            embeddings=embeddings,
            texts=texts,
            metadatas=[c["metadata"] for c in chunks],
        )
        await self._rebuild_bm25()
        log.info("domain_agent.ingest.done", domain=self.domain_key, chunks=len(chunks))
        return chunks

    async def retrieve(self, query: str) -> list[dict]:
        """Run hybrid search + rerank. Returns only chunks above min_rerank_score."""
        query_embedding = await self._embedding.embed(query)

        bm25_results = await asyncio.get_event_loop().run_in_executor(
            None, self._bm25.search, query, settings.bm25_top_k
        )
        vector_results = self._vector.search(query_embedding, settings.vector_top_k)

        fused_ids = rrf_merge(bm25_results, vector_results)
        candidates = self._vector.get_by_ids(fused_ids[:30])
        ranked = self._reranker.rerank(query, candidates, settings.reranker_top_n)

        # Drop chunks whose relevance score is below the configured threshold
        in_scope = [c for c in ranked if c.get("rerank_score", 0.0) >= settings.min_rerank_score]
        if len(in_scope) < len(ranked):
            log.info(
                "domain_agent.chunks_filtered",
                domain=self.domain_key,
                kept=len(in_scope),
                dropped=len(ranked) - len(in_scope),
            )
        return in_scope

    def build_context(self, ranked_chunks: list[dict]) -> str:
        return build_system_instruction(ranked_chunks)

    async def _rebuild_bm25(self) -> None:
        total = self._vector.total_chunks
        if total == 0:
            return
        results = self._vector._col.get(include=["documents"])
        self._bm25.rebuild(results["ids"], results["documents"])

    @property
    def chunk_count(self) -> int:
        return self._vector.total_chunks
