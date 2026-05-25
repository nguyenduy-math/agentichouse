from __future__ import annotations

import os
import unicodedata
from typing import Any


def _normalize(text: str) -> str:
    return unicodedata.normalize("NFC", text)


def _parse_pdf(file_path: str) -> list[dict[str, Any]]:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({"page": i, "text": _normalize(text)})
        if pages:
            return pages
    except Exception:
        pass

    try:
        import fitz
        doc = fitz.open(file_path)
        pages = []
        for i, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                pages.append({"page": i, "text": _normalize(text)})
        doc.close()
        return pages
    except Exception:
        return []


def _parse_docx(file_path: str) -> list[dict[str, Any]]:
    try:
        from docx import Document
        doc = Document(file_path)
        all_text: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text)
        pages: list[dict[str, Any]] = []
        current = ""
        page_num = 1
        for line in all_text:
            current += line + "\n"
            if len(current) >= 3000:
                pages.append({"page": page_num, "text": _normalize(current.strip())})
                page_num += 1
                current = ""
        if current.strip():
            pages.append({"page": page_num, "text": _normalize(current.strip())})
        return pages
    except Exception:
        return []


def _parse_text(file_path: str) -> list[dict[str, Any]]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        content = _normalize(content)
        if content.strip():
            return [{"page": 1, "text": content}]
        return []
    except Exception:
        return []


def parse_document(file_path: str) -> list[dict[str, Any]]:
    ext = os.path.splitext(file_path.lower())[1]
    if ext == ".pdf":
        return _parse_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    return _parse_text(file_path)
