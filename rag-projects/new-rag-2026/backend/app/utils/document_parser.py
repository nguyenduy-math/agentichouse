"""
Document parser: PDF (dual-library fallback) + DOCX (with table extraction).
All extracted text is NFC-normalized for Vietnamese diacritic consistency.

Carried over from graphrag-assistant with NFC normalization, dual PDF fallback,
and DOCX table extraction.
"""
from __future__ import annotations

import unicodedata
from pathlib import Path
from typing import Any


def _normalize(text: str) -> str:
    """Apply NFC normalization to ensure consistent Vietnamese diacritics."""
    return unicodedata.normalize("NFC", text)


def _parse_pdf(file_path: str) -> list[dict[str, Any]]:
    """
    Parse PDF to list of page dicts with 'page' and 'text' keys.
    Tries pdfplumber first (better for complex layouts), falls back to pymupdf.
    """
    pages: list[dict[str, Any]] = []

    # --- Attempt 1: pdfplumber (handles multi-column, tables well) ────────────
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                text = _normalize(text.strip())
                if text:
                    pages.append({"page": i, "text": text})
        if pages:
            return pages
    except Exception:
        pass

    # --- Attempt 2: pymupdf (more robust for scanned/protected PDFs) ──────────
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)
        for i, page in enumerate(doc, 1):
            text = page.get_text() or ""
            text = _normalize(text.strip())
            if text:
                pages.append({"page": i, "text": text})
        doc.close()
        if pages:
            return pages
    except Exception:
        pass

    return pages


def _parse_docx(file_path: str) -> list[dict[str, Any]]:
    """
    Parse DOCX extracting both paragraph text and table cell text.
    Returns a single page dict (DOCX has no page concept).
    """
    import docx  # python-docx

    doc = docx.Document(file_path)
    all_text: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            all_text.append(stripped)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped:
                    all_text.append(stripped)

    text = _normalize("\n".join(all_text))
    return [{"page": 1, "text": text}] if text else []


def _parse_txt(file_path: str) -> list[dict[str, Any]]:
    """Parse plain text file."""
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    text = _normalize(text.strip())
    return [{"page": 1, "text": text}] if text else []


def parse_document(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a document file and return a list of page dicts.
    Each dict has: {"page": int, "text": str}

    Supported formats: PDF, DOCX, TXT
    All text is NFC-normalized.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif suffix == ".txt":
        return _parse_txt(file_path)
    else:
        # Attempt plain text for unknown extensions
        try:
            return _parse_txt(file_path)
        except Exception:
            return []
